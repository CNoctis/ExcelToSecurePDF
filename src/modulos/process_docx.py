from docx import Document

# Función que lee nuestro archivo .docx
def read_docx(file_path: str) -> list:
    """Read the text from a docx file and return a list of paragraphs."""
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return text

# Función que creará una tabla para insertar en el archivo .docx pero reemplaza {tabla} por la tabla
def create_table_at_placeholder(doc, data:list, placeholder: str):
    """Create a table at the placeholder in a docx file."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph text at the placeholder
            before_text, after_text = paragraph.text.split(placeholder)
            
            # Clear the text in the current paragraph
            paragraph.text = before_text
            
            # Insert a new paragraph after the current one
            new_paragraph = paragraph.insert_paragraph_before("")
            
            # Add the table after the new paragraph
            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(data[0]):
                hdr_cells[i].text = header
            for row in data[1:]:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = cell
            
            # Move the table's XML to after the new paragraph
            new_paragraph._element.addnext(table._element)
            
            # Add the text after the placeholder back to the document
            if after_text:
                new_after_paragraph = doc.add_paragraph(after_text)
                table._element.addnext(new_after_paragraph._element)
            
            return doc
    
    # If the placeholder is not found, raise an exception
    raise ValueError(f"Placeholder '{placeholder}' not found in the document.")

# Función que remplaza {variable} por un valor en un archivo .docx
def replace_text(file_path: str, replacements: dict) -> Document:
    """Replace text in a docx file and return the Document object."""
    doc = Document(file_path)
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    return doc


# Función que guarda un archivo .docx
def save_docx(doc: Document, file_path: str) -> str:
    """Save a docx file and return the file path."""
    doc.save(file_path)
    return file_path
